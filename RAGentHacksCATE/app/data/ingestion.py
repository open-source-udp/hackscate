import os
from typing import List
from app.data.chunking import chunk_text
from app.data.chunking import __dict__ as _chunk_mod
from app.models.embeddings import EmbeddingClient
from app.utils import config
from app.utils.logger import logger
from math import sqrt

from langchain_chroma import Chroma
from langchain.schema import Document

from PyPDF2 import PdfReader
import docx


from app.data.marker import extract_text_with_marker

def read_pdf(path: str) -> str:
    reader = PdfReader(path)
    text = []
    for page in reader.pages:
        t = page.extract_text()
        text.append(t or "")
    _normalize = _chunk_mod.get('_normalize') if isinstance(_chunk_mod, dict) else None
    if callable(_normalize):
        return [_normalize(p) for p in text]
    else:
        return [p.strip() for p in text]

def read_docx(path: str) -> str:
    doc = docx.Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

def read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def load_file_to_text(path: str, use_marker_ocr: bool = True) -> str:
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        # Comportamiento condicional: primero extracción nativa; opcionalmente usar Marker según configuración y umbral
        raw_text = read_pdf(path)
        if use_marker_ocr and config.FORCE_MARKER_OCR:
            joined = "\n".join([p or "" for p in raw_text]) if isinstance(raw_text, list) else (raw_text or "")
            if len(joined) < config.MARKER_OCR_THRESHOLD:
                try:
                    ocr_result = extract_text_with_marker(path, force_ocr=True)
                    if isinstance(ocr_result, list):
                        return ocr_result
                    if '\f' in ocr_result:
                        return [p for p in ocr_result.split('\f')]
                    return [p.strip() for p in ocr_result.split('\n\n')]
                except Exception as e:
                    logger.exception(f"Marker OCR failed for {path}, falling back to PyPDF2: {e}")
        return raw_text

    elif ext in [".docx", ".doc"]:
        return read_docx(path)
    elif ext in [".txt", ".md"]:
        return read_txt(path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")

def ingest_files(paths: List[str], collection_name: str = None, persist: bool = None, dry_run: bool = False, dedup_threshold: float = None):
    def clean_text(text):
        return text.encode('utf-8', 'ignore').decode('utf-8')
    emb = EmbeddingClient()
    collection_name = collection_name or config.DEFAULT_COLLECTION_NAME
    vectordb = Chroma(
        persist_directory=config.CHROMA_PERSIST_DIR,
        embedding_function=emb._client,
        collection_name=collection_name,
    )
    persist = config.DEFAULT_PERSIST if persist is None else persist
    dedup_threshold = config.DEDUP_SIM_THRESHOLD if dedup_threshold is None else dedup_threshold
    documents = []
    seen_hashes = set()
    seen_embeddings = []
    for path in paths:
        # Usar OCR solo si está habilitado en configuración
        text = load_file_to_text(path, use_marker_ocr=config.FORCE_MARKER_OCR)
        # Validar texto
        if isinstance(text, list):
            combined_text = "\n".join([p or "" for p in text])
            if not combined_text.strip():
                logger.info(f"No text extracted from {path}, skipping.")
                continue
        else:
            if not text or len(text.strip()) == 0:
                logger.info(f"No text extracted from {path}, skipping.")
                continue

        # Un único chunk por archivo
        chunks = chunk_text(text, chunk_size_chars=config.CHUNK_DEFAULT_SIZE, chunk_overlap=config.CHUNK_DEFAULT_OVERLAP)
        for i, ch in enumerate(chunks):
            ch_dict = ch if isinstance(ch, dict) else {"text": ch}
            ch_text = ch_dict.get("text") or ""
            ch_clean = clean_text(ch_text)

            h = hash(ch_clean)
            if h in seen_hashes:
                logger.info(f"Skipping exact-duplicate chunk for {path} (chunk {i})")
                continue

            accept = True
            if dedup_threshold and dedup_threshold < 1.0 and len(seen_embeddings) > 0:
                emb_client = EmbeddingClient()
                q_emb = emb_client.embed([ch_clean])
                if isinstance(q_emb, list) and len(q_emb) == 1:
                    q_emb = q_emb[0]
                for d_emb in seen_embeddings:
                    dot = sum(a * b for a, b in zip(q_emb, d_emb))
                    norm_q = sqrt(sum(a * a for a in q_emb))
                    norm_d = sqrt(sum(a * a for a in d_emb))
                    sim = dot / (norm_q * norm_d) if norm_q and norm_d else 0.0
                    if sim >= dedup_threshold:
                        accept = False
                        logger.info(f"Skipping near-duplicate chunk for {path} (chunk {i}) sim={sim:.3f})")
                        break

            if not accept:
                continue

            seen_hashes.add(h)
            emb_client = EmbeddingClient()
            emb = emb_client.embed([ch_clean])
            if isinstance(emb, list) and len(emb) == 1:
                seen_embeddings.append(emb[0])

            metadata = {"source": os.path.basename(path), "chunk": i}
            if ch_dict.get("page_start") is not None:
                metadata.update({"page_start": ch_dict.get("page_start"), "page_end": ch_dict.get("page_end")})
            if ch_dict.get("char_start") is not None:
                metadata.update({"char_start": ch_dict.get("char_start"), "char_end": ch_dict.get("char_end")})
            documents.append(Document(page_content=ch_clean, metadata=metadata))

    if documents:
        if dry_run:
            logger.info(f"Dry-run ingest: {len(documents)} documents would be added to collection '{collection_name}'")
        else:
            vectordb.add_documents(documents)
            logger.info(f"Ingested {len(documents)} documents into collection '{collection_name}'")
    else:
        logger.info("No documents to add after processing (dedup/filter may have removed all chunks)")
    return documents


def ingest_file(path: str, collection_name: str, persist: bool = None, dry_run: bool = False, dedup_threshold: float = None):
        """
        Ingesta un único archivo en la colección indicada.
        Parámetros:
            - path: ruta_de_archivo
            - collection_name: nombre de la colección destino
            - persist, dry_run, dedup_threshold: igual que ingest_files
        """
        return ingest_files([path], collection_name=collection_name, persist=persist, dry_run=dry_run, dedup_threshold=dedup_threshold)
